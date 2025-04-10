import pandas as pd
import os
import re
from docx import Document
from lotsize import lot_size_dict  # Import the dictionary from the lotsize.py file

folder_path = r'E:\Stocks\Options csv'
output_file = 'Affordable_Options.docx'

# Create a new Word document
doc = Document()
doc.add_heading('Affordable Stock Options', level=1)

# Loop through all CSV files in the folder
for filename in os.listdir(folder_path):
    if filename.endswith('.csv'):
        file_path = os.path.join(folder_path, filename)

        # Extract stock symbol using regex
        match = re.search(r'option-chain-ED-(.*?)-\d{1,2}-[A-Za-z]+-\d{4}', filename)
        symbol = match.group(1) if match else 'UNKNOWN'

        # Get lot size (default to 1000 if not found)
        lot_size = lot_size_dict.get(symbol, 1000)

        # Read CSV, skipping the first row
        df = pd.read_csv(file_path, skiprows=1)
        df.columns = df.columns.str.upper()

        # Rename necessary columns
        df = df.rename(columns={
            df.columns[5]: 'CALL LTP',
            df.columns[3]: 'CALL VOLUME',
            df.columns[17]: 'PUT LTP',
            df.columns[19]: 'PUT VOLUME',
            'STRIKE': 'STRIKE PRICE'
        })

        # Convert necessary columns to numeric
        df[['CALL LTP', 'PUT LTP', 'CALL VOLUME', 'PUT VOLUME']] = df[
            ['CALL LTP', 'PUT LTP', 'CALL VOLUME', 'PUT VOLUME']
        ].apply(pd.to_numeric, errors='coerce')

        # Filter affordable options
        affordable_calls = df[df['CALL LTP'] * lot_size <= 2500][['STRIKE PRICE', 'CALL LTP', 'CALL VOLUME']]
        affordable_puts = df[df['PUT LTP'] * lot_size <= 2500][['STRIKE PRICE', 'PUT LTP', 'PUT VOLUME']]

        if not affordable_calls.empty or not affordable_puts.empty:
            doc.add_heading(symbol, level=2)  # Stock name as a heading

        # Add Calls Table
        if not affordable_calls.empty:
            doc.add_paragraph("Affordable CALL Options:")
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'

            # Add headers
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Strike Price'
            hdr_cells[1].text = 'Call LTP'
            hdr_cells[2].text = 'Call Volume'

            # Add data
            for _, row in affordable_calls.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(row['STRIKE PRICE'])
                row_cells[1].text = str(row['CALL LTP'])
                row_cells[2].text = str(row['CALL VOLUME'])

        # Add Puts Table
        if not affordable_puts.empty:
            doc.add_paragraph("Affordable PUT Options:")
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'

            # Add headers
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Strike Price'
            hdr_cells[1].text = 'Put LTP'
            hdr_cells[2].text = 'Put Volume'

            # Add data
            for _, row in affordable_puts.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(row['STRIKE PRICE'])
                row_cells[1].text = str(row['PUT LTP'])
                row_cells[2].text = str(row['PUT VOLUME'])

        doc.add_paragraph("\n")  # Add space between stocks

# Save the Word document
doc.save(output_file)
print(f"Data has been saved to '{output_file}'")
